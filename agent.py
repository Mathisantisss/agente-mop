"""
Agente conversacional experto en el Manual de Carreteras de Chile (MOP).
Uso:
    python agent.py
"""

import json
import os
import re
import subprocess
import sys
import unicodedata
from datetime import datetime
from pathlib import Path

import anthropic
import chromadb
import openpyxl
from custom_embeddings import FastEmbedFunction
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from rich.console import Console
from rich.markdown import Markdown
from rich.panel import Panel
from rich.prompt import Prompt
from rich.rule import Rule
from rich.table import Table
from rich.text import Text

from mop_config import (
    ANTHROPIC_MODEL,
    CHROMA_DIR,
    COLLECTION_NAME,
    EMBEDDING_MODEL,
    REPORTS_DIR,
    TOP_K,
)

console = Console()

# ---------------------------------------------------------------------------
# ChromaDB — inicialización lazy
# ---------------------------------------------------------------------------
_collection = None


def get_collection():
    global _collection
    if _collection is None:
        embedding_fn = FastEmbedFunction(model_name=EMBEDDING_MODEL)
        client = chromadb.PersistentClient(path=str(CHROMA_DIR))
        _collection = client.get_or_create_collection(
            name=COLLECTION_NAME,
            embedding_function=embedding_fn,
            metadata={"hnsw:space": "cosine"},
        )
    return _collection


# ---------------------------------------------------------------------------
# Herramientas
# ---------------------------------------------------------------------------

def buscar_en_manual(consulta: str, volumen: str = "", n_resultados: int = TOP_K) -> str:
    """Búsqueda semántica en ChromaDB."""
    collection = get_collection()

    if collection.count() == 0:
        return "⚠ La base de datos está vacía. Ejecuta primero: python ingest.py"

    where = {"volumen": {"$eq": volumen}} if volumen else None

    try:
        results = collection.query(
            query_texts=[consulta],
            n_results=min(n_resultados, collection.count()),
            where=where,
            include=["documents", "metadatas", "distances"],
        )
    except Exception as e:
        return f"Error al consultar la base de datos: {e}"

    docs = results["documents"][0]
    metas = results["metadatas"][0]
    dists = results["distances"][0]

    if not docs:
        return "No se encontraron resultados relevantes."

    MAX_CHARS_POR_FRAGMENTO = 600  # limitar tokens enviados a Claude

    lines = [f"**Resultados para:** '{consulta}'\n"]
    for i, (doc, meta, dist) in enumerate(zip(docs, metas, dists), 1):
        relevance = max(0, (1 - dist) * 100)
        lines.append(
            f"### Resultado {i} — {meta.get('volumen', 'N/D')} "
            f"(págs. {meta.get('pagina_inicio', '?')}–{meta.get('pagina_fin', '?')}) "
            f"| Relevancia: {relevance:.1f}%\n"
        )
        texto = doc.strip()
        if len(texto) > MAX_CHARS_POR_FRAGMENTO:
            texto = texto[:MAX_CHARS_POR_FRAGMENTO] + "..."
        lines.append(texto)
        lines.append("")

    return "\n".join(lines)


def listar_volumenes() -> str:
    """Lista los volúmenes disponibles en la base de datos."""
    collection = get_collection()

    if collection.count() == 0:
        return "⚠ La base de datos está vacía. Ejecuta primero: python ingest.py"

    all_meta = collection.get(include=["metadatas"])["metadatas"]
    volumes: dict[str, int] = {}
    for meta in all_meta:
        vol = meta.get("volumen", "Sin clasificar")
        volumes[vol] = volumes.get(vol, 0) + 1

    lines = ["**Volúmenes disponibles en la base de datos:**\n"]
    for vol, count in sorted(volumes.items()):
        lines.append(f"- **{vol}**: {count} fragmentos")

    lines.append(f"\n**Total de fragmentos:** {sum(volumes.values())}")
    return "\n".join(lines)


def _safe_filename(name: str) -> str:
    """Convierte un nombre a nombre de archivo seguro."""
    nfkd = unicodedata.normalize("NFKD", name)
    ascii_str = nfkd.encode("ascii", "ignore").decode("ascii")
    safe = re.sub(r"[^\w\s\-]", "", ascii_str).strip()
    safe = re.sub(r"\s+", "_", safe)
    return safe[:80] or "archivo"


def generar_excel(datos: str, nombre_archivo: str, titulo: str = "", hoja: str = "Datos") -> str:
    """Genera un archivo Excel (.xlsx) con los datos proporcionados."""
    try:
        parsed = json.loads(datos)
    except json.JSONDecodeError as e:
        return f"Error: el parámetro 'datos' debe ser un JSON válido. Detalle: {e}"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = hoja[:31]

    header_fill = PatternFill("solid", fgColor="1F497D")
    header_font = Font(bold=True, color="FFFFFF", name="Calibri", size=11)
    alt_fill = PatternFill("solid", fgColor="DCE6F1")
    border_side = Side(style="thin", color="B8CCE4")
    cell_border = Border(
        left=border_side, right=border_side, top=border_side, bottom=border_side
    )

    row_offset = 1

    # Título opcional
    if titulo:
        ws.cell(row=1, column=1, value=titulo).font = Font(bold=True, size=13, name="Calibri")
        row_offset = 3

    def style_header(row: int, ncols: int):
        for c in range(1, ncols + 1):
            cell = ws.cell(row=row, column=c)
            cell.fill = header_fill
            cell.font = header_font
            cell.border = cell_border
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def style_data(row: int, ncols: int, alt: bool):
        for c in range(1, ncols + 1):
            cell = ws.cell(row=row, column=c)
            if alt:
                cell.fill = alt_fill
            cell.border = cell_border
            cell.alignment = Alignment(vertical="center", wrap_text=True)

    # Caso 1: lista de dicts
    if isinstance(parsed, list) and parsed and isinstance(parsed[0], dict):
        headers = list(parsed[0].keys())
        for c, h in enumerate(headers, 1):
            ws.cell(row=row_offset, column=c, value=h)
        style_header(row_offset, len(headers))

        for i, row_data in enumerate(parsed):
            r = row_offset + 1 + i
            for c, h in enumerate(headers, 1):
                ws.cell(row=r, column=c, value=row_data.get(h, ""))
            style_data(r, len(headers), i % 2 == 1)

    # Caso 2: lista de listas
    elif isinstance(parsed, list) and parsed and isinstance(parsed[0], list):
        headers = parsed[0]
        for c, h in enumerate(headers, 1):
            ws.cell(row=row_offset, column=c, value=h)
        style_header(row_offset, len(headers))

        for i, row_data in enumerate(parsed[1:]):
            r = row_offset + 1 + i
            for c, val in enumerate(row_data, 1):
                ws.cell(row=r, column=c, value=val)
            style_data(r, len(row_data), i % 2 == 1)

    # Caso 3: dict de secciones
    elif isinstance(parsed, dict):
        current_row = row_offset
        for section_name, section_data in parsed.items():
            ws.cell(row=current_row, column=1, value=section_name).font = Font(
                bold=True, size=12, name="Calibri", color="1F497D"
            )
            current_row += 1

            if isinstance(section_data, list) and section_data:
                if isinstance(section_data[0], dict):
                    headers = list(section_data[0].keys())
                    for c, h in enumerate(headers, 1):
                        ws.cell(row=current_row, column=c, value=h)
                    style_header(current_row, len(headers))
                    current_row += 1

                    for i, row_data in enumerate(section_data):
                        for c, h in enumerate(headers, 1):
                            ws.cell(row=current_row, column=c, value=row_data.get(h, ""))
                        style_data(current_row, len(headers), i % 2 == 1)
                        current_row += 1
                else:
                    for item in section_data:
                        ws.cell(row=current_row, column=1, value=str(item))
                        current_row += 1

            current_row += 1

    # Auto-ajustar columnas
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                max_len = max(max_len, len(str(cell.value or "")))
            except Exception:
                pass
        ws.column_dimensions[col_letter].width = min(max(max_len + 4, 12), 50)

    safe_name = _safe_filename(nombre_archivo)
    output_path = REPORTS_DIR / f"{safe_name}.xlsx"
    wb.save(output_path)

    return f"Excel generado correctamente: {output_path}"


def generar_reporte(
    contenido: str,
    nombre_archivo: str,
    titulo: str = "Reporte",
    subtitulo: str = "",
) -> str:
    """Genera un reporte Word (.docx) con formato profesional."""
    doc = Document()

    # Estilos de fuente base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Encabezado con fecha
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"Manual de Carreteras MOP — {datetime.now().strftime('%d/%m/%Y')}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.runs[0].font.size = Pt(9)
    header_para.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # Título principal
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(titulo)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    if subtitulo:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(subtitulo)
        sub_run.font.size = Pt(13)
        sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()

    # Parsear Markdown → docx
    for line in contenido.splitlines():
        stripped = line.strip()

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        elif stripped.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            p.runs[0].font.size = Pt(11)

        elif stripped == "---" or stripped == "___":
            doc.add_paragraph("─" * 60)

        elif stripped == "":
            doc.add_paragraph()

        else:
            # Parsear negrita inline **texto**
            p = doc.add_paragraph()
            p.runs  # ensure initialized
            parts = re.split(r"\*\*(.+?)\*\*", stripped)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(11)
                if j % 2 == 1:  # era negrita
                    run.bold = True

    safe_name = _safe_filename(nombre_archivo)
    output_path = REPORTS_DIR / f"{safe_name}.docx"
    doc.save(output_path)

    return f"Reporte Word generado correctamente: {output_path}"


def generar_pdf(
    contenido: str,
    nombre_archivo: str,
    titulo: str = "Reporte",
    subtitulo: str = "",
) -> str:
    """Genera un reporte PDF (.pdf) con formato profesional usando reportlab.

    Acepta el mismo formato Markdown que generar_reporte:
    # / ## / ### para titulos, - para listas, **negrita**.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.lib.colors import HexColor
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        from reportlab.platypus import (
            Paragraph, SimpleDocTemplate, Spacer, PageBreak, HRFlowable,
        )
    except ImportError:
        return ("Error: reportlab no esta instalado. Ejecuta: "
                "pip install reportlab")

    safe_name = _safe_filename(nombre_archivo)
    output_path = REPORTS_DIR / f"{safe_name}.pdf"

    # Documento A4 con margenes profesionales
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        topMargin=2.2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        title=titulo,
        author="Agente MOP",
    )

    azul_mop = HexColor("#1F497D")
    gris = HexColor("#404040")

    base_styles = getSampleStyleSheet()
    estilo_titulo = ParagraphStyle(
        "TituloMOP", parent=base_styles["Title"],
        fontName="Helvetica-Bold", fontSize=18,
        textColor=azul_mop, alignment=TA_CENTER, spaceAfter=8,
    )
    estilo_subtitulo = ParagraphStyle(
        "SubtituloMOP", parent=base_styles["Normal"],
        fontName="Helvetica", fontSize=12,
        textColor=gris, alignment=TA_CENTER, spaceAfter=16,
    )
    estilo_h1 = ParagraphStyle(
        "H1MOP", parent=base_styles["Heading1"],
        fontName="Helvetica-Bold", fontSize=14,
        textColor=azul_mop, spaceBefore=14, spaceAfter=8,
    )
    estilo_h2 = ParagraphStyle(
        "H2MOP", parent=base_styles["Heading2"],
        fontName="Helvetica-Bold", fontSize=12,
        textColor=azul_mop, spaceBefore=12, spaceAfter=6,
    )
    estilo_h3 = ParagraphStyle(
        "H3MOP", parent=base_styles["Heading3"],
        fontName="Helvetica-Bold", fontSize=11,
        textColor=azul_mop, spaceBefore=10, spaceAfter=4,
    )
    estilo_parrafo = ParagraphStyle(
        "ParrafoMOP", parent=base_styles["Normal"],
        fontName="Helvetica", fontSize=10.5,
        leading=15, alignment=TA_JUSTIFY, spaceAfter=6,
    )
    estilo_bullet = ParagraphStyle(
        "BulletMOP", parent=estilo_parrafo,
        leftIndent=18, bulletIndent=6, spaceAfter=3,
    )

    def _md_inline_a_html(linea: str) -> str:
        """Convierte **negrita** y `code` a tags HTML que reportlab entiende."""
        # Escapar caracteres especiales de reportlab
        linea = linea.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        linea = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", linea)
        linea = re.sub(r"`(.+?)`", r'<font name="Courier" color="#C0392B">\1</font>', linea)
        return linea

    elementos = []

    # Cabecera
    elementos.append(Paragraph(titulo, estilo_titulo))
    if subtitulo:
        elementos.append(Paragraph(subtitulo, estilo_subtitulo))
    elementos.append(Paragraph(
        f"Manual de Carreteras MOP &nbsp;·&nbsp; {datetime.now().strftime('%d/%m/%Y')}",
        ParagraphStyle("MetaMOP", parent=estilo_parrafo,
                       fontSize=8.5, textColor=gris, alignment=TA_CENTER),
    ))
    elementos.append(Spacer(1, 0.4 * cm))
    elementos.append(HRFlowable(width="100%", thickness=1, color=azul_mop, spaceAfter=10))

    # Cuerpo
    for linea in contenido.splitlines():
        s = linea.rstrip()
        if not s.strip():
            elementos.append(Spacer(1, 0.18 * cm))
        elif s.startswith("### "):
            elementos.append(Paragraph(_md_inline_a_html(s[4:]), estilo_h3))
        elif s.startswith("## "):
            elementos.append(Paragraph(_md_inline_a_html(s[3:]), estilo_h2))
        elif s.startswith("# "):
            elementos.append(Paragraph(_md_inline_a_html(s[2:]), estilo_h1))
        elif s.lstrip().startswith(("- ", "* ", "• ")):
            texto = s.lstrip()[2:]
            elementos.append(Paragraph(
                _md_inline_a_html(texto), estilo_bullet, bulletText="•",
            ))
        elif s.strip() in ("---", "___", "***"):
            elementos.append(HRFlowable(width="100%", thickness=0.5,
                                        color=gris, spaceBefore=4, spaceAfter=4))
        else:
            elementos.append(Paragraph(_md_inline_a_html(s), estilo_parrafo))

    # Footer con numero de pagina
    def _footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(gris)
        canvas.drawRightString(
            doc.pagesize[0] - 2.2 * cm, 1.2 * cm,
            f"Pagina {doc.page}",
        )
        canvas.drawString(
            2.2 * cm, 1.2 * cm,
            "Ministerio de Obras Publicas - Chile",
        )
        canvas.restoreState()

    doc.build(elementos, onFirstPage=_footer, onLaterPages=_footer)

    return f"PDF generado correctamente: {output_path}"


# ---------------------------------------------------------------------------
# Schemas de herramientas para la API
# ---------------------------------------------------------------------------

TOOLS = [
    {
        "name": "listar_codigos_senales",
        "description": (
            "Retorna el LISTADO COMPLETO Y EXACTO de todos los códigos de señales "
            "del Manual MOP, extraídos de forma determinística (precisión 100%). "
            "Usar SIEMPRE que el usuario pregunte por: tipos de señales, códigos, "
            "nomenclatura, listado de señalética, tabla de señales. "
            "Acepta filtro por categoría: RP (reglamentaria prioridad), RPO (prohibición), "
            "RR (restricción), PI (intersección), PO (operativas), IS (servicios), "
            "ID (dirección), IP (preseñalización), D (demarcaciones), TM/TP/TR (trabajos)."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "categoria": {
                    "type": "string",
                    "description": "Filtro opcional por categoría (ej: 'IS', 'RP'). Vacío = todas.",
                },
            },
            "required": [],
        },
    },
    {
        "name": "buscar_en_manual",
        "description": (
            "Busca información técnica en los volúmenes del Manual de Carreteras de Chile (MOP). "
            "Usa esta herramienta para responder preguntas sobre normas, diseño vial, pavimentos, "
            "estructuras, señalización, seguridad vial, especificaciones técnicas y cualquier tema "
            "del manual. Devuelve fragmentos relevantes con su volumen y páginas de origen."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "consulta": {
                    "type": "string",
                    "description": "Pregunta o términos técnicos a buscar en el manual.",
                },
                "volumen": {
                    "type": "string",
                    "description": (
                        "Filtrar por volumen específico, ej: 'Volumen 3'. "
                        "Dejar vacío para buscar en todos los volúmenes."
                    ),
                },
                "n_resultados": {
                    "type": "integer",
                    "description": "Número de fragmentos a recuperar (por defecto 5, máximo 15).",
                    "default": 5,
                },
            },
            "required": ["consulta"],
        },
    },
    {
        "name": "listar_volumenes",
        "description": (
            "Lista todos los volúmenes del Manual de Carreteras disponibles en la base de datos, "
            "junto con el número de fragmentos indexados de cada uno."
        ),
        "input_schema": {
            "type": "object",
            "properties": {},
            "required": [],
        },
    },
    {
        "name": "generar_excel",
        "description": (
            "Genera un archivo Excel (.xlsx) con datos tabulares. Usar cuando el usuario solicite "
            "una tabla, planilla, comparativo, listado o cualquier salida en formato Excel. "
            "Los datos deben ser un JSON válido en una de estas formas:\n"
            "1. Lista de dicts: [{\"col1\": val, \"col2\": val}, ...]\n"
            "2. Lista de listas (primera fila = encabezados): [[\"Col1\",\"Col2\"],[v1,v2],...]\n"
            "3. Dict de secciones: {\"Sección A\": [{...}], \"Sección B\": [{...}]}"
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "datos": {
                    "type": "string",
                    "description": "Datos en formato JSON (ver descripción).",
                },
                "nombre_archivo": {
                    "type": "string",
                    "description": "Nombre del archivo sin extensión, ej: 'normas_velocidad_diseno'.",
                },
                "titulo": {
                    "type": "string",
                    "description": "Título opcional que aparece en la parte superior de la hoja.",
                },
                "hoja": {
                    "type": "string",
                    "description": "Nombre de la hoja del Excel (por defecto 'Datos').",
                },
            },
            "required": ["datos", "nombre_archivo"],
        },
    },
    {
        "name": "generar_pdf",
        "description": (
            "Genera un reporte profesional en PDF (.pdf). Usar cuando el usuario solicite "
            "explícitamente un PDF (no un Word). Ideal para informes finales, documentos "
            "de presentación o entregables que no requieren edición. El contenido puede "
            "incluir formato Markdown básico: # ## ### para títulos, - para listas, "
            "**texto** para negrita, `código` para código."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "contenido": {
                    "type": "string",
                    "description": "Contenido del PDF en formato Markdown.",
                },
                "nombre_archivo": {
                    "type": "string",
                    "description": "Nombre del archivo sin extensión, ej: 'informe_pavimentos'.",
                },
                "titulo": {
                    "type": "string",
                    "description": "Título principal del documento.",
                },
                "subtitulo": {
                    "type": "string",
                    "description": "Subtítulo o descripción breve (opcional).",
                },
            },
            "required": ["contenido", "nombre_archivo"],
        },
    },
    {
        "name": "generar_reporte",
        "description": (
            "Genera un reporte profesional en Word (.docx). Usar cuando el usuario solicite "
            "un informe, reporte, documento técnico o memoria EN WORD/EDITABLE. El contenido "
            "puede incluir formato Markdown básico: # ## ### para títulos, - para listas, "
            "**texto** para negrita."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "contenido": {
                    "type": "string",
                    "description": "Contenido del reporte en formato Markdown.",
                },
                "nombre_archivo": {
                    "type": "string",
                    "description": "Nombre del archivo sin extensión, ej: 'informe_pavimentos'.",
                },
                "titulo": {
                    "type": "string",
                    "description": "Título principal del documento.",
                },
                "subtitulo": {
                    "type": "string",
                    "description": "Subtítulo o descripción breve (opcional).",
                },
            },
            "required": ["contenido", "nombre_archivo"],
        },
    },
]

def listar_codigos_senales(categoria: str = "") -> str:
    """Wrapper sobre signs_db para que sea llamable como herramienta."""
    try:
        from signs_db import formatear_listado_completo, formatear_listado_por_categoria
        if categoria:
            return formatear_listado_por_categoria(categoria)
        return formatear_listado_completo()
    except Exception as e:
        return f"Error al obtener señales: {e}"


TOOL_FUNCTIONS = {
    "listar_codigos_senales": listar_codigos_senales,
    "buscar_en_manual": buscar_en_manual,
    "listar_volumenes": listar_volumenes,
    "generar_excel": generar_excel,
    "generar_reporte": generar_reporte,
    "generar_pdf": generar_pdf,
}


def dispatch_tool(name: str, inputs: dict) -> str:
    fn = TOOL_FUNCTIONS.get(name)
    if fn is None:
        return f"Error: herramienta '{name}' no encontrada."
    try:
        return str(fn(**inputs))
    except Exception as e:
        return f"Error al ejecutar '{name}': {e}"


# ---------------------------------------------------------------------------
# Prompt de sistema
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = """Eres un experto en el Manual de Carreteras de Chile, publicado por el \
Ministerio de Obras Públicas (MOP). Tienes acceso a los 9 volúmenes completos del manual \
a través de una base de datos vectorial.

**Tus responsabilidades:**
- Responder preguntas técnicas sobre diseño vial, pavimentos, estructuras, señalización, \
seguridad vial, hidrología, puentes, túneles y todo lo contenido en el Manual MOP.
- Citar siempre el volumen y las páginas de referencia al dar información técnica.
- Usar la herramienta `buscar_en_manual` para recuperar información del manual antes de responder.
- Generar archivos Excel cuando el usuario solicite tablas, comparativos o planillas.
- Generar reportes Word cuando el usuario solicite informes, memorias o documentos técnicos.

**Cómo responder:**
- Sé preciso y técnico, pero claro y accesible.
- Estructura tus respuestas con secciones cuando sea apropiado.
- Si no encuentras información en el manual para una consulta específica, indícalo claramente.
- Responde siempre en español.
- Al citar el manual, usa el formato: (Manual MOP, Volumen X, pág. Y).

**Restricciones:**
- Solo responde sobre temas del Manual de Carreteras MOP o temas directamente relacionados \
con ingeniería vial chilena.
- No inventes normas ni valores técnicos; si no están en el manual, indícalo."""


# ---------------------------------------------------------------------------
# Utilidades de sesión
# ---------------------------------------------------------------------------

def abrir_archivo(path: Path) -> None:
    """Abre un archivo con la aplicación por defecto del sistema."""
    try:
        if sys.platform == "win32":
            os.startfile(str(path))
        elif sys.platform == "darwin":
            subprocess.run(["open", str(path)])
        else:
            subprocess.run(["xdg-open", str(path)])
    except Exception:
        pass


def mostrar_ayuda() -> None:
    table = Table(title="Comandos disponibles", border_style="cyan", show_header=True)
    table.add_column("Comando", style="bold yellow")
    table.add_column("Descripción")
    table.add_row("/ayuda", "Muestra este menú de comandos")
    table.add_row("/volumenes", "Lista los volúmenes indexados en la base de datos")
    table.add_row("/reportes", "Abre la carpeta de reportes generados")
    table.add_row("/limpiar", "Nueva sesión (borra el historial de conversación)")
    table.add_row("/salir", "Termina el agente")
    console.print(table)


def mostrar_bienvenida() -> None:
    console.print(
        Panel(
            "[bold cyan]Agente MOP — Manual de Carreteras de Chile[/bold cyan]\n"
            "[dim]Experto en los 9 volúmenes del Manual MOP · Jun. 2025[/dim]\n\n"
            "[white]Consultas tecnicas, comparativos y generacion de reportes Excel/Word.[/white]\n"
            "[dim]Escribe [bold]/ayuda[/bold] para ver los comandos disponibles.[/dim]",
            title="[bold white]MOP — Ministerio de Obras Publicas[/bold white]",
            border_style="cyan",
            padding=(1, 3),
        )
    )


def verificar_base_datos() -> int:
    """Retorna el numero de fragmentos en la base. Muestra aviso si esta vacia."""
    try:
        client = chromadb.PersistentClient(path=str(CHROMA_DIR))
        col = client.get_collection(COLLECTION_NAME)
        return col.count()
    except Exception:
        return 0


# ---------------------------------------------------------------------------
# Loop principal del agente
# ---------------------------------------------------------------------------

def run_agent() -> None:
    client = anthropic.Anthropic()
    messages: list[dict] = []

    mostrar_bienvenida()

    # Verificar base de datos al inicio
    n_frags = verificar_base_datos()
    if n_frags == 0:
        console.print(
            Panel(
                "[yellow]La base de datos esta vacia.[/yellow]\n"
                "Ejecuta primero: [bold]python ingest.py[/bold]\n"
                "El agente puede funcionar, pero no tendra acceso al contenido del manual.",
                border_style="yellow",
            )
        )
    else:
        console.print(f"[dim]Base de datos lista: {n_frags:,} fragmentos indexados.[/dim]\n")

    while True:
        try:
            user_input = Prompt.ask("\n[bold cyan]Tu[/bold cyan]").strip()
        except (KeyboardInterrupt, EOFError):
            console.print("\n[dim]Sesion terminada.[/dim]")
            break

        if not user_input:
            continue

        # --- Comandos de sesión ---
        cmd = user_input.lower()

        if cmd in ("/salir", "/exit", "/quit"):
            console.print("[dim]Hasta luego.[/dim]")
            break

        if cmd in ("/limpiar", "/clear", "/nuevo"):
            messages = []
            console.print("[dim]Sesion limpiada. Nueva conversacion iniciada.[/dim]")
            continue

        if cmd in ("/ayuda", "/help"):
            mostrar_ayuda()
            continue

        if cmd in ("/volumenes", "/vols"):
            with console.status("Consultando volumenes...", spinner="dots"):
                resultado = listar_volumenes()
            console.print(Panel(Markdown(resultado), border_style="blue", title="Volumenes MOP"))
            continue

        if cmd in ("/reportes", "/archivos"):
            console.print(f"[dim]Abriendo carpeta: {REPORTS_DIR}[/dim]")
            abrir_archivo(REPORTS_DIR)
            continue

        # --- Consulta al agente ---
        messages.append({"role": "user", "content": user_input})
        archivos_generados: list[Path] = []

        with console.status(
            "[bold yellow]Consultando el Manual MOP...[/bold yellow]", spinner="dots"
        ):
            # Loop interno: ejecutar herramientas hasta end_turn
            while True:
                response = client.messages.create(
                    model=ANTHROPIC_MODEL,
                    max_tokens=8192,
                    system=SYSTEM_PROMPT,
                    tools=TOOLS,
                    messages=messages,
                )

                text_blocks = []
                tool_uses = []

                for block in response.content:
                    if block.type == "text":
                        text_blocks.append(block.text)
                    elif block.type == "tool_use":
                        tool_uses.append(block)

                messages.append({"role": "assistant", "content": response.content})

                if response.stop_reason == "end_turn" or not tool_uses:
                    break

                # Ejecutar herramientas
                tool_results = []
                for tool_block in tool_uses:
                    result = dispatch_tool(tool_block.name, tool_block.input)

                    # Detectar archivos generados para abrirlos después
                    if "generado correctamente:" in result:
                        try:
                            path_str = result.split("generado correctamente:")[-1].strip()
                            p = Path(path_str)
                            if p.exists():
                                archivos_generados.append(p)
                        except Exception:
                            pass

                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": tool_block.id,
                        "content": result,
                    })

                messages.append({"role": "user", "content": tool_results})

        # --- Mostrar respuesta ---
        final_text = "\n".join(text_blocks).strip()
        if final_text:
            console.print(
                Panel(
                    Markdown(final_text),
                    title="[bold green]Agente MOP[/bold green]",
                    border_style="green",
                    padding=(1, 2),
                )
            )

        # Abrir archivos generados automáticamente
        for archivo in archivos_generados:
            console.print(f"[dim]Abriendo: {archivo.name}[/dim]")
            abrir_archivo(archivo)


if __name__ == "__main__":
    run_agent()
